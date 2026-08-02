"""Resume text parsing utilities.

Supports parsing resume text from:
- Raw text input (pasted resume content)
- PDF files (using pdfplumber)
- DOCX files (using python-docx)
"""

import re
from io import BytesIO
from pathlib import Path


class ResumeValidationError(ValueError):
    """A validation failure that is safe to display directly to the user."""


_ALLOWED_PDF_MIME_TYPES = {"application/pdf", "application/x-pdf"}
_ALLOWED_DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/octet-stream",
}
_MAX_PDF_BYTES = 10 * 1024 * 1024
_MAX_DOCX_BYTES = 10 * 1024 * 1024
_MAX_TEXT_CHARACTERS = 100_000
_MIN_READABLE_CHARACTERS = 80


_RESUME_SECTIONS = {
    "experience": re.compile(
        r"\b(?:work\s+experience|professional\s+experience|employment|internships?)\b",
        re.IGNORECASE,
    ),
    "education": re.compile(
        r"\b(?:education|academic\s+(?:background|qualifications?)|qualifications?)\b",
        re.IGNORECASE,
    ),
    "skills": re.compile(
        r"\b(?:technical\s+skills|skills|technologies|competencies)\b",
        re.IGNORECASE,
    ),
    "projects": re.compile(r"\bprojects?\b", re.IGNORECASE),
    "certifications": re.compile(r"\bcertifications?\b", re.IGNORECASE),
    "objective": re.compile(r"\b(?:career\s+)?objective\b", re.IGNORECASE),
    "summary": re.compile(
        r"\b(?:professional\s+|career\s+|profile\s+)?summary\b",
        re.IGNORECASE,
    ),
    "achievements": re.compile(
        r"\b(?:achievements?|awards?|accomplishments?)\b",
        re.IGNORECASE,
    ),
}

_CONTACT_PATTERNS = (
    re.compile(r"\b[\w.+-]+@[\w.-]+\.[a-z]{2,}\b", re.IGNORECASE),
    re.compile(r"(?<!\d)(?:\+?\d[\s().-]*){10,15}(?!\d)"),
)

_NON_RESUME_PATTERNS = (
    re.compile(
        r"\b(?:certificate\s+of\s+(?:completion|achievement|participation)|"
        r"this\s+is\s+to\s+certify)\b",
        re.IGNORECASE,
    ),
    re.compile(
        r"\b(?:invoice\s*(?:number|no\.?|#)|bill\s+to|payment\s+due|"
        r"subtotal|total\s+amount|tax\s+invoice|gstin)\b",
        re.IGNORECASE,
    ),
    re.compile(
        r"\b(?:isbn(?:-1[03])?|table\s+of\s+contents|chapter\s+\d+|"
        r"lecture\s+notes|class\s+notes|study\s+notes)\b",
        re.IGNORECASE,
    ),
    re.compile(
        r"\b(?:survey|questionnaire|respondent|feedback\s+form|"
        r"strongly\s+agree|strongly\s+disagree|select\s+one|tick\s+one)\b",
        re.IGNORECASE,
    ),
)

_SECTION_HEADING_WORDS = {
    "resume",
    "curriculum vitae",
    "education",
    "skills",
    "experience",
    "projects",
    "certifications",
    "objective",
    "summary",
    "achievements",
}


def detect_candidate_name(text: str) -> str | None:
    """Return a likely candidate name from the first few non-empty lines."""
    for line in [item.strip() for item in text.splitlines() if item.strip()][:10]:
        normalized = line.lower().rstrip(":")
        if normalized in _SECTION_HEADING_WORDS:
            continue
        if any(pattern.search(line) for pattern in _CONTACT_PATTERNS):
            continue
        if len(line) > 60 or any(char.isdigit() for char in line):
            continue
        if re.fullmatch(
            r"[A-Za-z][A-Za-z.'-]+(?:\s+[A-Za-z][A-Za-z.'-]+){1,3}",
            line,
        ):
            return line
    return None


def assess_resume_text(text: str) -> tuple[bool, list[str]]:
    """Check for multiple structural signals before treating text as a resume."""
    normalized = " ".join((text or "").split())
    if len(normalized) < 120:
        return False, ["The document has too little readable resume information."]

    section_hits = [
        name for name, pattern in _RESUME_SECTIONS.items() if pattern.search(normalized)
    ]
    contact_hits = sum(bool(pattern.search(normalized)) for pattern in _CONTACT_PATTERNS)
    chronology = bool(
        re.search(
            r"\b(?:19|20)\d{2}\b|\b\d+\+?\s+(?:years?|months?)\b|\b(?:present|current)\b",
            normalized,
            re.IGNORECASE,
        )
    )
    role_signal = bool(
        re.search(
            r"\b(?:developer|engineer|analyst|manager|designer|consultant|intern|accountant|teacher)\b",
            normalized,
            re.IGNORECASE,
        )
    )
    non_resume_hits = sum(
        bool(pattern.search(normalized)) for pattern in _NON_RESUME_PATTERNS
    )

    reasons = []
    if len(section_hits) < 2:
        reasons.append(
            "Include resume sections such as Experience, Education, Skills, or Projects."
        )
    if contact_hits == 0:
        reasons.append("No valid email address or phone number was found.")
    if not chronology and not role_signal:
        reasons.append("No work-role, date, or experience details were found.")
    if non_resume_hits:
        reasons.append(
            "The PDF appears to be a certificate, invoice, bill, book, notes, "
            "or another non-resume document."
        )

    is_resume = (
        len(section_hits) >= 2
        and contact_hits >= 1
        and (chronology or role_signal)
        and non_resume_hits == 0
    )
    return is_resume, reasons


def validate_resume_text(text: str) -> str:
    """Return cleaned resume text or raise a user-safe validation error."""
    cleaned = parse_resume_text(text)
    if len(cleaned) > _MAX_TEXT_CHARACTERS:
        raise ResumeValidationError(
            "The pasted resume is too long. Please keep it under 100,000 characters."
        )
    accepted, reasons = assess_resume_text(cleaned)
    if not accepted:
        detail = " ".join(reasons)
        raise ResumeValidationError(
            "This document does not appear to be a resume/CV. "
            "A resume must include an email and/or phone number and at least two "
            "standard sections such as Education, Skills, Experience, Projects, "
            f"Certifications, Objective, Summary, or Achievements. {detail}"
        )
    return cleaned


def _read_uploaded_file(file) -> tuple[bytes, str | None, str | None]:
    """Read an upload without trusting its extension or browser metadata."""
    if isinstance(file, (str, Path)):
        path = Path(file)
        data = path.read_bytes()
        return data, path.name, None

    filename = getattr(file, "name", None)
    mime_type = getattr(file, "type", None)
    if hasattr(file, "getvalue"):
        data = file.getvalue()
    else:
        position = file.tell() if hasattr(file, "tell") else None
        data = file.read()
        if position is not None and hasattr(file, "seek"):
            file.seek(position)
    return bytes(data), filename, mime_type


def _validate_pdf_container(
    data: bytes, filename: str | None, mime_type: str | None
) -> None:
    """Validate extension, MIME metadata, size, and the real PDF signature."""
    if filename and Path(filename).suffix.lower() != ".pdf":
        raise ResumeValidationError(
            "Only PDF files are accepted. Files such as JPG, PNG, DOC, DOCX, "
            "and TXT cannot be uploaded."
        )
    if mime_type and mime_type.lower() not in _ALLOWED_PDF_MIME_TYPES:
        raise ResumeValidationError(
            f"The uploaded file reports the type '{mime_type}', not PDF. "
            "Please upload an original PDF resume."
        )
    if not data:
        raise ResumeValidationError("The uploaded PDF is empty.")
    if len(data) > _MAX_PDF_BYTES:
        raise ResumeValidationError(
            "The PDF is larger than 10 MB. Please upload a smaller resume PDF."
        )
    # PDF signatures normally begin at byte zero, but ISO 32000 readers permit
    # a small amount of leading data. Checking the first 1 KB is strict enough
    # to reject renamed images/documents while remaining standards-friendly.
    if b"%PDF-" not in data[:1024]:
        raise ResumeValidationError(
            "This is not a real PDF file. Renaming another file to '.pdf' is "
            "not supported."
        )


def _extract_pdf_text(data: bytes) -> tuple[str, bool]:
    """Extract readable text and report whether image content was encountered."""
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError(
            "PDF extraction is unavailable. Install pdfplumber and try again."
        ) from exc

    text_parts = []
    contains_images = False
    try:
        with pdfplumber.open(BytesIO(data)) as pdf:
            if not pdf.pages:
                raise ResumeValidationError("The uploaded PDF has no pages.")
            for page in pdf.pages:
                contains_images = contains_images or bool(page.images)
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(page_text)
    except ResumeValidationError:
        raise
    except Exception as exc:
        raise ResumeValidationError(
            "The PDF is corrupt, password-protected, or cannot be read. "
            "Please export the original resume as a new PDF and try again."
        ) from exc
    return "\n".join(text_parts), contains_images


def validate_resume_pdf_upload(file) -> str:
    """Validate a genuine, text-readable resume PDF and return its text.

    All checks run server-side. The ML pipeline must only receive the returned
    text; any validation failure raises ``ResumeValidationError`` with a
    user-friendly explanation.
    """
    data, filename, mime_type = _read_uploaded_file(file)
    _validate_pdf_container(data, filename, mime_type)
    text, contains_images = _extract_pdf_text(data)
    readable = " ".join(text.split())
    if len(readable) < _MIN_READABLE_CHARACTERS:
        if contains_images:
            raise ResumeValidationError(
                "No readable text was found. The PDF appears to be scanned or "
                "image-only; please upload a text-based PDF resume."
            )
        raise ResumeValidationError(
            "The PDF is blank or contains too little readable text to be a resume."
        )
    return validate_resume_text(text)


def validate_resume_docx_upload(file) -> str:
    """Validate a DOCX upload, extract its text, and confirm it is a resume."""
    data, filename, mime_type = _read_uploaded_file(file)
    if filename and Path(filename).suffix.lower() != ".docx":
        raise ResumeValidationError("Only genuine DOCX files are accepted here.")
    if mime_type and mime_type.lower() not in _ALLOWED_DOCX_MIME_TYPES:
        raise ResumeValidationError(
            f"The uploaded file reports the type '{mime_type}', not DOCX."
        )
    if not data:
        raise ResumeValidationError("The uploaded DOCX is empty.")
    if len(data) > _MAX_DOCX_BYTES:
        raise ResumeValidationError(
            "The DOCX is larger than 10 MB. Please upload a smaller resume."
        )
    if not data.startswith(b"PK"):
        raise ResumeValidationError(
            "This is not a real DOCX file. Renaming another file is not supported."
        )

    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError(
            "DOCX extraction is unavailable. Install python-docx and try again."
        ) from exc

    try:
        document = Document(BytesIO(data))
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            paragraphs.extend(cell.text for row in table.rows for cell in row.cells)
    except Exception as exc:
        raise ResumeValidationError(
            "The DOCX is corrupt, password-protected, or cannot be read."
        ) from exc

    readable = "\n".join(item for item in paragraphs if item.strip())
    if len(" ".join(readable.split())) < _MIN_READABLE_CHARACTERS:
        raise ResumeValidationError(
            "The DOCX is blank or contains too little readable text to be a resume."
        )
    return validate_resume_text(readable)


def validate_resume_upload(file) -> str:
    """Dispatch a supported PDF or DOCX upload to its secure validator."""
    filename = str(getattr(file, "name", "") or "")
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return validate_resume_pdf_upload(file)
    if suffix == ".docx":
        return validate_resume_docx_upload(file)
    raise ResumeValidationError("Upload a PDF or DOCX resume only.")


def parse_resume_text(text: str) -> str:
    """Clean and normalize pasted resume text.

    Parameters
    ----------
    text : str
        Raw resume text pasted by the user.

    Returns
    -------
    str
        Cleaned resume text ready for feature extraction.
    """
    if not text or not isinstance(text, str):
        return ""

    # Remove common resume formatting artifacts
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Remove multiple blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Remove leading/trailing whitespace
    text = text.strip()

    return text


def parse_resume_pdf(file) -> str:
    """Extract text from a PDF resume file using pdfplumber.

    Parameters
    ----------
    file : file-like or str
        A file path (str) or a file-like object (e.g., Streamlit UploadedFile)
        pointing to a PDF file.

    Returns
    -------
    str
        Extracted text from the PDF.
    """
    data, _, _ = _read_uploaded_file(file)
    text, _ = _extract_pdf_text(data)
    return text


def extract_skills_from_text(text: str) -> list:
    """Extract potential skill keywords from resume text.

    Uses a curated list of common technical and soft skills
    to identify skills mentioned in the resume.

    Parameters
    ----------
    text : str
        Resume text (raw or cleaned).

    Returns
    -------
    list of str
        Sorted list of detected skills.
    """
    if not text:
        return []

    text_lower = text.lower()

    # Comprehensive skill keywords list
    skill_keywords = sorted(set([
        # Programming Languages
        "python", "java", "javascript", "c++", "c#", "r", "sql", "php",
        "ruby", "go", "scala", "swift", "kotlin", "typescript", "matlab",
        "sas", "perl", "rust", "haskell", "lua",
        # Web Technologies
        "html", "css", "react", "angular", "vue", "node.js", "django",
        "flask", "fastapi", "spring", "express", "bootstrap", "jquery",
        "rest api", "graphql",
        # Data Science & ML
        "machine learning", "deep learning", "natural language processing",
        "nlp", "computer vision", "tensorflow", "keras", "pytorch",
        "scikit-learn", "sklearn", "xgboost", "pandas", "numpy",
        "scipy", "matplotlib", "seaborn", "plotly", "tableau",
        "data analysis", "data visualization", "statistical modeling",
        "regression", "classification", "clustering", "neural network",
        "convolutional neural network", "cnn", "recurrent neural network",
        "rnn", "lstm", "transformer", "bert", "gpt",
        # Data Engineering
        "hadoop", "spark", "hive", "kafka", "airflow", "etl",
        "data pipeline", "big data", "sql", "nosql", "mongodb",
        "postgresql", "mysql", "oracle", "cassandra", "redis",
        "elasticsearch", "databricks", "snowflake", "redshift",
        # Cloud & DevOps
        "aws", "azure", "gcp", "google cloud", "docker", "kubernetes",
        "jenkins", "git", "github", "gitlab", "ci/cd", "terraform",
        "ansible", "linux", "bash", "shell scripting",
        # Business & Management
        "project management", "agile", "scrum", "jira", "leadership",
        "communication", "team management", "strategic planning",
        "business analysis", "requirements gathering", "stakeholder management",
        "budgeting", "forecasting",
        # Accounting & Finance
        "accounting", "financial analysis", "budgeting", "forecasting",
        "accounts payable", "accounts receivable", "general ledger",
        "tax preparation", "audit", "compliance", "gaap", "ifrs",
        "quickbooks", "sap", "erp",
        # Engineering
        "cad", "autocad", "solidworks", "matlab", "simulink",
        "circuit design", "pcb design", "embedded systems",
        "iot", "robotics", "plc", "scada",
        # Soft Skills
        "problem solving", "critical thinking", "teamwork",
        "time management", "attention to detail", "adaptability",
        "creativity", "analytical thinking",
        # Marketing
        "digital marketing", "seo", "sem", "social media marketing",
        "content marketing", "email marketing", "google analytics",
        "crm", "salesforce", "market research",
    ]))

    detected = []
    for skill in skill_keywords:
        # Token boundaries prevent false positives such as the language "r"
        # matching every ordinary word that contains that letter.
        pattern = rf"(?<!\w){re.escape(skill)}(?!\w)"
        if re.search(pattern, text_lower):
            detected.append(skill)

    return sorted(set(detected))
